"""
Author: Renhuan Huang
This script will iterate through the files inside a given folder
for each word document, the key words (contained in {} sign will be replaced by values read from csv file)

"""
from docx import Document
from csv import reader

import os
import sys

# preparation of files

absolute_path = r"C:\Python\Projects and Scripts\docx"
parameter_file = "doc parameter PP.csv"
counter = 0

# end of file preparation

# To read document id from a csv file, and store it into a list.

list_of_docid = []    # Create an empty list for storage of parameters read from csv file.
list_of_parameter = ['{System Desc}', '{System ID}', '{TS ID}', '{Dept}', '{System Impact}', '{System Brand}', '{System Model}', '{IOQ ID}', '{Doc ID}']  # Parameter placeholder inside Word document

with open(parameter_file,'r') as read_obj:
    csv_reader = reader(read_obj,)
    list_of_docid = list(csv_reader) # this is a two dimensional list, row is first dimension, columns are second.

#print(list_of_docid[1][0])

# loop thru the list for file names.

for file_name in os.listdir(absolute_path):
        
    if file_name.find('P19000') == -1:    # Need to make sure you only work on your Word file in the folder. No need to work on sub-folder or other file type.
        continue
    
    # To open the first file for auto-population of key words
    document = Document(file_name)    # Create an instance of document object for further handling
    para = document.paragraphs      # paragraphs inside contents
    #header = document.headers       # paragraph inside header
    #table = document.tables         # paragraph inside table

   # To find the row corresponding to the file name
    parameter_for_given_file_name = []    # this list will hold the parameter matching the file name. it is an element of the list_of_docid
    key_word_in_file_name = file_name[0:17]
   # print(type(key_word_in_file_name))

    for row in list_of_docid:
       if row[15] == key_word_in_file_name:
           parameter_for_given_file_name = row
           #print(file_name, parameter_for_given_file_name[15])


   #

   # Do text replacement, by iterating through paragraphs in 1) contents; 2) headers; 3) tables
   
   # 1) Iterate through contents
    for p in para:
        for place_holder in list_of_parameter:
            if place_holder == '{System Desc}': para_index = 0
            if place_holder == '{System ID}': para_index = 18
            if place_holder == '{TS ID}': para_index = 13
            if place_holder == '{Dept}': para_index = 20
            if place_holder == '{System Impact}': para_index = 10
            if place_holder == '{System Brand}': para_index = 2
            if place_holder == '{System Model}': para_index = 2
            if place_holder == '{IOQ ID}': para_index = 14
            if place_holder == '{Doc ID}': para_index = 15
            #style = p.style
            p.text = p.text.replace(place_holder,parameter_for_given_file_name[para_index])
            #p.stype = style

    # 2) iterate through headers

    for section in document.sections:
        
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for cell_para in cell.paragraphs:
                            for place_holder in list_of_parameter:
                                if place_holder == '{System Desc}': para_index = 0
                                if place_holder == '{Doc ID}': para_index = 15
             #                   style = cell_para.style
                                cell_para.text = cell_para.text.replace(place_holder,parameter_for_given_file_name[para_index])
              #                  cell_para.style = style

    # 3) iterate through tables 
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                for cell_para in cell.paragraphs:
                    for place_holder in list_of_parameter:
                     if place_holder == '{System Desc}': para_index = 0
                     if place_holder == '{System ID}': para_index = 18
                     if place_holder == '{TS ID}': para_index = 13
                     if place_holder == '{Dept}': para_index = 20
                     if place_holder == '{System Impact}': para_index = 10
                     if place_holder == '{System Brand}': para_index = 1
                     if place_holder == '{System Model}': para_index = 2
                     if place_holder == '{IOQ ID}': para_index = 14
                     if place_holder == '{Doc ID}': para_index = 15
#                     style = p.style
                     cell_para.text = cell_para.text.replace(place_holder,parameter_for_given_file_name[para_index])
 #                    p.style = style
        
    document.save(file_name)    # Save all the changes back into the orginal file.

