"""
Author: Renhuan Huang

Do not run this unless you know what you're doing.

This script will iterate through the files inside a given folder which contains both RTM files and TS files
The csv file contains the mapping between TS ID and RTM ID. Now what does the script do:
    - Iterate throught the folder, and look for RTM file
    - Then, extract the RTM id from file name
    - Then use the RTM id, to search the corresponding TS ID inside the csv file
    - Once TS id is found, go back to iterate through folder for TS file name
    - Then assign the RTM file name and TS file name into destination_file and source_file

    - Then open the RTM and TS file
    - Then copy the specifications from TS into TRM and save it.

    - The total time spent in copying is approximiately 1 minute, for 58 RTM files.

Now how to use:
    - Copy all RTM document and TS document inside the same folder
    - Copy the csv file containing TS ID and RTM ID inside the same folder
    - Run the script inside Python
"""
from docx import Document

from csv import reader
import os
# preparation of files

absolute_path = r"C:\Python\Projects and Scripts\docx\TS to RTM"
parameter_file = "TS RTM Mapping.csv"
counter = 0

# end of file preparation

# To read document id from a csv file, and store it into a list.

list_of_docid_mapping = []    # Create an empty list for storage of parameters read from csv file.

with open(parameter_file,'r') as read_obj:
    csv_reader = reader(read_obj,)
    list_of_docid_mapping = list(csv_reader) # this is a two dimensional list, row is first dimension, columns are second.

##

# Iterate thru the list for file names.
file_names = os.listdir(absolute_path)
source_file_name=""
destination_file_name = ""

for file_name in file_names:
        
    if file_name.find('P19000-QV-RTM') == -1:    # Need to make sure you only work on your Word file in the folder. No need to work on sub-folder or other file type. in the current case, only look for RTM document.
        continue
    counter += 1
    
    destination_file_name = file_name

    rtm_id = file_name[0:17]
    ts_id = ""
    
    for r in list_of_docid_mapping:
        if rtm_id == r[1]:
            ts_id = r[0]

    source_file_name = ""
    for ts_doc in file_names:
        if ts_id in ts_doc:
            source_file_name = ts_doc
           
    source_doc = Document(source_file_name)
    destination_doc = Document(destination_file_name)

    destination_tables = destination_doc.tables
    destination_RTM_table= destination_tables[len(destination_tables)-1]

    # Iterate through all tables inside source_doc. because the numbered text can not exracted, therefore they ahve to be manually added into the destination file. 
    # values in Column 2 and 3 will be copied.

    source_tables = source_doc.tables
    d_row_count = 1

    for tb in source_tables:
        d_row_count += 1
        if tb.cell(0,0).text != 'Identifier': 
            d_row_count -= 1
            continue  # exclude tables that are not specification related.
        for rows in range(1,len(tb.rows)):
            d_row_count += 1
            for columns in range(1,len(tb.columns)):
                #print(tb.cell(rows,columns).text)
                
                #print(d_row_count)
                if len(destination_RTM_table.rows) < d_row_count + 5:
                    for i in range(5):
                        #print(len(destination_RTM_table.rows))
                        destination_RTM_table.add_row()
                        #print(len(destination_RTM_table.rows))
                destination_RTM_table.cell(d_row_count,columns).text = tb.cell(rows,columns).text
        
        
    #print(len(destination_RTM_table.rows))
    destination_doc.save(destination_file_name)
    print("OK with no errors found")
print(counter, " RTM files were processed")